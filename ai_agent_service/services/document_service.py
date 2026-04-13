from docx import Document
import os

class DocumentService:
    @staticmethod
    def generate_docx(content: str, filename: str, title: str = "BÁO CÁO TỰ ĐỘNG - NCKH_ZMA") -> str:
        if not os.path.exists("static"):
            os.makedirs("static")
            
        doc = Document()
        doc.add_heading(title, 0)
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        filepath = f"static/{filename}"
        doc.save(filepath)
        return filepath
